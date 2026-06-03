"""
Gerador do Relatório Analítico de Área — CompStat Municipal
Lê input JSON (area + synthesis) e gera .docx no formato CompStat.
"""
import argparse
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            tag.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def heading_para(doc, text: str, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9 if level == 1 else 8)
    run.font.color.rgb = RGBColor(0xE8, 0x5D, 0x26)
    return p


def body_para(doc, text: str, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.bold = bold
    return p


def small_para(doc, text: str, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(8)
    if color:
        run.font.color.rgb = RGBColor(*[int(color[i:i+2], 16) for i in (0, 2, 4)])
    return p


def hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), '2a2a30')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─── Report Builder ──────────────────────────────────────────────────────────

def build_report(area: dict, synthesis: str, output_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    stats = area['stats']
    score = area['score']
    nome = area['nome']
    hoje = date.today().strftime('%d/%m/%Y')

    # ── CABEÇALHO ────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('RELATÓRIO ANALÍTICO DE ÁREA — COMPSTAT MUNICIPAL')
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xE8, 0x5D, 0x26)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    rs = subtitle.add_run('Subsídio para Reunião de CompStat | Prefeitura do Rio de Janeiro')
    rs.font.size = Pt(8)
    rs.font.color.rgb = RGBColor(0x70, 0x70, 0x78)

    # Area + date header table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    set_cell_bg(tbl.cell(0, 0), '111113')
    set_cell_bg(tbl.cell(0, 1), '111113')
    c0 = tbl.cell(0, 0).paragraphs[0]
    c0.add_run('Área FM').font.size = Pt(8)
    c0.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x78)
    tbl.cell(0, 0).add_paragraph().add_run(nome).font.size = Pt(10)
    c1 = tbl.cell(0, 1).paragraphs[0]
    c1.add_run('Período').font.size = Pt(8)
    c1.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x78)
    p1 = tbl.cell(0, 1).add_paragraph()
    p1.add_run(f"Crimes (ISP-RJ): 2020–2024\nChamados 1746: 2020–2024\nFatores campo: 2026 · Denúncias DD: 2025").font.size = Pt(9)

    doc.add_paragraph()

    # ── RESUMO EXECUTIVO ─────────────────────────────────────────────────────
    heading_para(doc, '0. Resumo Executivo')

    # 4 norteadoras
    perguntas = [
        ('Locais de maior incidência coincidem com a rota da FM?',
         f"Sim. Top trechos: {', '.join(t['locf_norm'].title() for t in area['top_trechos'][:3])}. Score de risco: {score['total']:.0f}/100."),
        ('Horário de maior incidência coincide com QMD?',
         f"Pico em {stats['pico_horario']}. Cobertura deverá priorizar janela noturna 19h–23h."),
        ('Dinâmica criminal coincide com o modelo de emprego da FM?',
         synthesis[:200] + '…' if synthesis else 'Aguardando síntese qualitativa.'),
        ('Fatores urbanos relevantes estão sendo resolvidos?',
         f"{stats['fatores_urbanos_total']} fatores mapeados na área, {len(area['fatores_por_orgao'])} órgãos responsáveis."),
    ]

    tbl2 = doc.add_table(rows=len(perguntas) + 1, cols=3)
    tbl2.style = 'Table Grid'
    hdrs = ['Pergunta Norteadora', 'Diagnóstico com base nos dados', 'Observações']
    for j, h in enumerate(hdrs):
        c = tbl2.cell(0, j)
        set_cell_bg(c, '18181b')
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x70, 0x70, 0x78)
    for i, (q, a) in enumerate(perguntas):
        row = i + 1
        tbl2.cell(row, 0).paragraphs[0].add_run(q).font.size = Pt(8)
        tbl2.cell(row, 1).paragraphs[0].add_run(a).font.size = Pt(8)
        tbl2.cell(row, 2).paragraphs[0].add_run('—').font.size = Pt(8)

    doc.add_paragraph()

    # ── 1. OCORRÊNCIAS ───────────────────────────────────────────────────────
    heading_para(doc, '1. Ocorrências Criminais')

    # Identificação
    heading_para(doc, 'Identificação da Área', level=2)
    id_tbl = doc.add_table(rows=3, cols=4)
    id_tbl.style = 'Table Grid'
    ch1746 = area.get('chamados_1746', {})
    ch_total = ch1746.get('total', 0)
    ch_pct_at = ch1746.get('pct_atendido', 0)

    id_data = [
        [('Área FM', nome), ('Câmeras', str(stats['cameras_total'])),
         ('Score de Risco', f"{score['total']:.0f}/100"), ('Data', hoje)],
        [('Crimes (ISP-RJ)', str(stats['crimes_total'])),
         ('Pico Horário', stats['pico_horario']),
         ('DD (denúncia crime)', str(stats.get('denuncias_total', 0))),
         ('Fatores Campo', str(stats['fatores_urbanos_total']))],
        [('1746 (serviço público)', f"{ch_total:,}".replace(',', '.')),
         ('1746 % Atendidos', f"{ch_pct_at}%"),
         ('PSR', str(stats.get('psr_total', 0))),
         ('', '')],
    ]
    for r_i, row_data in enumerate(id_data):
        set_cell_bg(id_tbl.cell(r_i, 0), '111113')
        for c_i, (label, val) in enumerate(row_data):
            cell = id_tbl.cell(r_i, c_i)
            lp = cell.paragraphs[0]
            lr = lp.add_run(label)
            lr.font.size = Pt(8)
            lr.font.color.rgb = RGBColor(0x70, 0x70, 0x78)
            vp = cell.add_paragraph()
            vr = vp.add_run(val)
            vr.font.size = Pt(10)
            vr.bold = True

    doc.add_paragraph()

    # Indicadores
    heading_para(doc, 'Distribuição por Tipo de Ocorrência', level=2)
    tipo_tbl = doc.add_table(rows=1, cols=3)
    tipo_tbl.style = 'Table Grid'
    for j, h in enumerate(['Tipo', 'Quantidade', '% do Total']):
        c = tipo_tbl.cell(0, j)
        set_cell_bg(c, '18181b')
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x70, 0x70, 0x78)

    total = stats['crimes_total'] or 1
    for tipo, count in sorted(stats['crimes_por_tipo'].items(), key=lambda x: -x[1]):
        row = tipo_tbl.add_row()
        row.cells[0].paragraphs[0].add_run(tipo).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(str(count)).font.size = Pt(9)
        row.cells[2].paragraphs[0].add_run(f"{count/total*100:.1f}%").font.size = Pt(9)

    doc.add_paragraph()

    # Análise temporal
    heading_para(doc, 'Análise Temporal', level=2)
    h_dist = stats.get('hora_distribution', {})
    if h_dist:
        pico_h = max(h_dist, key=h_dist.get)
        pico_v = h_dist[pico_h]
        body_para(doc, f"Pico de incidência: {pico_h}h com {pico_v} ocorrências.")

    d_dist = stats.get('dia_distribution', {})
    if d_dist:
        pico_dia = max(d_dist, key=d_dist.get)
        body_para(doc, f"Dia da semana mais crítico: {pico_dia} ({d_dist[pico_dia]} ocorrências).")

    doc.add_paragraph()

    # ── 2. DINÂMICA CRIMINAL ──────────────────────────────────────────────────
    heading_para(doc, '2. Dinâmica Criminal')
    if synthesis:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(synthesis)
        r.font.size = Pt(9)
    else:
        body_para(doc, '[Síntese qualitativa não disponível. Execute a síntese com IA na plataforma.]')

    doc.add_paragraph()

    # ── 3. EFETIVO EMPREGADO ─────────────────────────────────────────────────
    heading_para(doc, '3. Efetivo Empregado — Força Municipal')
    fm_tbl = doc.add_table(rows=1, cols=4)
    fm_tbl.style = 'Table Grid'
    for j, h in enumerate(['Campo', 'Situação Atual', 'Sugestão de Alteração', 'Justificativa']):
        c = fm_tbl.cell(0, j)
        set_cell_bg(c, '18181b')
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x70, 0x70, 0x78)

    fm_rows = [
        ('Nº de Agentes / Turno', 'N/D', 'N/D', 'Definir na reunião'),
        ('Horário de Cobertura', 'N/D', f"{stats['pico_horario']} (pico)"),
        ('Modelo de Emprego', 'N/D', 'Definir com base na dinâmica'),
        ('Locais de Cobertura', ', '.join(t['locf_norm'].title() for t in area['top_trechos'][:2]), 'Manter foco nos trechos críticos'),
    ]
    for row_data in fm_rows:
        row = fm_tbl.add_row()
        for j, val in enumerate(row_data):
            row.cells[j].paragraphs[0].add_run(val).font.size = Pt(9)

    doc.add_paragraph()

    # ── 4. FATORES URBANOS ────────────────────────────────────────────────────
    heading_para(doc, '4. Fatores de Incidência Criminal')
    small_para(doc,
        'Fontes distintas: "Campo" = observação da equipe FM (fatores urbanos). '
        '"1746" = chamados da Central de Atendimento ao cidadão (serviços públicos). '
        '"DD" = Disque Denúncia (crime anônimo). NÃO confundir 1746 com DD.',
        '707078')

    if area['fatores_por_orgao']:
        fu_tbl = doc.add_table(rows=1, cols=3)
        fu_tbl.style = 'Table Grid'
        for j, h in enumerate(['Fator Identificado', 'Descrição / Incidência', 'Responsável']):
            c = fu_tbl.cell(0, j)
            set_cell_bg(c, '18181b')
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x70, 0x70, 0x78)
        for fo in area['fatores_por_orgao']:
            for t in fo['tipos'][:3]:
                row = fu_tbl.add_row()
                row.cells[0].paragraphs[0].add_run(t['tipo']).font.size = Pt(9)
                row.cells[1].paragraphs[0].add_run(f"{t['count']} registros na área").font.size = Pt(9)
                row.cells[2].paragraphs[0].add_run(fo['orgao']).font.size = Pt(9)

    doc.add_paragraph()

    # ── 4.1 VALIDAÇÃO CRUZADA ────────────────────────────────────────────────
    validacao = area.get('validacao_cruzada', [])
    if validacao:
        heading_para(doc, '4.1 Validação Cruzada — Campo × Cidadão (1746)')
        small_para(doc,
            'Compara fatores observados em campo (equipe FM) com chamados 1746 (cidadão). '
            'Quando ambos concordam, o problema é crônico e validado.',
            '707078')

        vc_tbl = doc.add_table(rows=1, cols=5)
        vc_tbl.style = 'Table Grid'
        for j, h in enumerate(['Órgão', 'Campo (fatores)', '1746 (chamados)', '% Atendidos', 'Vencidos']):
            c = vc_tbl.cell(0, j)
            set_cell_bg(c, '18181b')
            r = c.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(0x70, 0x70, 0x78)
        for v in validacao:
            row = vc_tbl.add_row()
            row.cells[0].paragraphs[0].add_run(v['orgao']).font.size = Pt(9)
            row.cells[1].paragraphs[0].add_run(str(v.get('fatores_campo', 0))).font.size = Pt(9)
            row.cells[2].paragraphs[0].add_run(str(v.get('chamados_1746', 0))).font.size = Pt(9)
            pct = round(v['chamados_atendidos'] / max(v['chamados_1746'], 1) * 100) if v.get('chamados_1746') else 0
            row.cells[3].paragraphs[0].add_run(f"{pct}%").font.size = Pt(9)
            row.cells[4].paragraphs[0].add_run(str(v.get('chamados_vencidos', 0))).font.size = Pt(9)

        doc.add_paragraph()

    # ── 5. PLANO DE AÇÃO ──────────────────────────────────────────────────────
    heading_para(doc, '5. Plano de Ação e Responsabilização')
    small_para(doc, 'Preencher durante/após a reunião CompStat.', '707078')

    pa_tbl = doc.add_table(rows=1, cols=4)
    pa_tbl.style = 'Table Grid'
    for j, h in enumerate(['Ação Acordada', 'Responsável', 'Prazo', 'Status']):
        c = pa_tbl.cell(0, j)
        set_cell_bg(c, '18181b')
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x70, 0x70, 0x78)

    # Pre-populate from fatores
    for fo in area['fatores_por_orgao'][:4]:
        if fo['tipos']:
            top_tipo = fo['tipos'][0]['tipo']
            row = pa_tbl.add_row()
            row.cells[0].paragraphs[0].add_run(f"Resolver: {top_tipo}").font.size = Pt(9)
            row.cells[1].paragraphs[0].add_run(fo['orgao']).font.size = Pt(9)
            row.cells[2].paragraphs[0].add_run('A definir').font.size = Pt(9)
            row.cells[3].paragraphs[0].add_run('—').font.size = Pt(9)

    # Add blank rows for filling
    for _ in range(3):
        row = pa_tbl.add_row()
        for cell in row.cells:
            cell.paragraphs[0].add_run(' ')

    doc.add_paragraph()
    small_para(doc, f'Gerado automaticamente pela Plataforma CompStat Rio em {hoje}. Validar e editar antes de usar.', '44444c')

    doc.save(output_path)
    print(f'Relatório salvo: {output_path}')


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--input', required=True)
    parser.add_argument('--output', required=True)
    args = parser.parse_args()

    with open(args.input, encoding='utf-8') as f:
        data = json.load(f)

    build_report(data['area'], data.get('synthesis', ''), args.output)


if __name__ == '__main__':
    main()
