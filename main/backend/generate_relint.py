"""
Gerador de RELINT — replica EXATAMENTE o formato RI_010–RI_017.
Estrutura:
  1. Tabela header 3 colunas: FORÇA MUNICIPAL | RESERVADO | Pág.
  2. Parágrafo bold fora da tabela: RELATÓRIO DE INTELIGÊNCIA DE ÁREA – COMPSTAT – DADOS PÚBLICOS
  3. Tabela principal 1 coluna sem bordas:
       Row 0: RELATÓRIO DE INTELIGÊNCIA DE ÁREA / Subsídio para Reunião de CompStat
       Row 1: NOME DA ÁREA
       Row 2: Introdução
       Row N (par): TÍTULO DA SUB-ÁREA
       Row N+1 (ímpar): conteúdo com bullets + dinâmica
       ...
       Row -2: CONCLUSÃO
       Row -1: texto conclusão + bullets + fechamento
"""
import argparse, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ──────────────────────────────────────────────────────────────────

def no_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('left','top','right','bottom','insideH','insideV'):
                el = OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), 'none')
                el.set(qn('w:sz'), '0')
                tcBorders.append(el)
            tcPr.append(tcBorders)


def run(para, text, bold=False, size=10, color=None):
    r = para.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Arial'
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r


def cell_para(cell, text, bold=False, size=10, space_before=0, space_after=2):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run(p, text, bold=bold, size=size)
    return p


def add_row(table, lines):
    """lines = list of (text, bold)"""
    row  = table.add_row()
    cell = row.cells[0]
    # remove default empty para
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for text, bold in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3 if bold else 0)
        p.paragraph_format.space_after  = Pt(3 if bold else 2)
        run(p, text, bold=bold, size=10)
    return row


# ── main builder ─────────────────────────────────────────────────────────────

def build_relint(data: dict, output_path: str):
    area    = data['area']
    content = data['content']

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2)
    sec.right_margin  = Cm(2)
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)

    # remove default paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # ── 1. Header table 3 colunas ─────────────────────────────────────────
    hdr = doc.add_table(rows=1, cols=3)
    hdr.style = 'Table Grid'
    # column widths: ~5cm | ~8cm | ~5cm
    widths = [Cm(4.5), Cm(9), Cm(4.5)]
    for i, w in enumerate(widths):
        hdr.columns[i].width = w

    cells = hdr.rows[0].cells
    for c, txt in zip(cells, ['FORÇA MUNICIPAL', 'RESERVADO', 'Pág.']):
        p = c.paragraphs[0]
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(9)
        r.font.name = 'Arial'

    # ── 2. Título fora da tabela ─────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_before = Pt(6)
    title_para.paragraph_format.space_after  = Pt(6)
    run(title_para, 'RELATÓRIO ', bold=True, size=10)
    run(title_para, 'DE INTELIGÊNCIA', bold=True, size=10)
    run(title_para, ' DE ÁREA – COMPSTAT – DADOS PÚBLICOS', bold=True, size=10)

    # ── 3. Tabela principal (1 coluna, sem bordas) ────────────────────────
    tbl = doc.add_table(rows=0, cols=1)
    tbl.style = 'Table Grid'

    # Row 0: cabeçalho do relatório
    add_row(tbl, [
        ('RELATÓRIO DE INTELIGÊNCIA DE ÁREA', True),
        ('Subsídio para Reunião de CompStat', True),
    ])

    # Row 1: nome da área
    add_row(tbl, [(area['nome'].upper(), True)])

    # Row 2: introdução
    add_row(tbl, [(content['intro'], False)])

    # Rows pares/ímpares: sub-áreas
    for sec_data in content['secoes']:
        # título
        add_row(tbl, [(sec_data['titulo'].upper(), True)])

        # conteúdo
        lines = [(sec_data['texto'], False)]
        if sec_data.get('bullets'):
            lines.append(('Também foram identificados:', False))
            for b in sec_data['bullets']:
                lines.append((f'• {b}', False))
        if sec_data.get('dinamica'):
            lines.append((sec_data['dinamica'], False))
        add_row(tbl, lines)

    # Conclusão
    add_row(tbl, [('CONCLUSÃO', True)])

    conc = content['conclusao']
    conc_lines = [(conc['texto'], False)]
    if conc.get('intro_bullets'):
        conc_lines.append((conc['intro_bullets'], False))
    for b in conc.get('bullets', []):
        conc_lines.append((f'• {b}', False))
    if conc.get('fechamento'):
        conc_lines.append((conc['fechamento'], False))
    add_row(tbl, conc_lines)

    no_borders(tbl)
    doc.save(output_path)
    print(f'RELINT salvo: {output_path}')


def main():
    p = argparse.ArgumentParser()
    p.add_argument('--input',  required=True)
    p.add_argument('--output', required=True)
    args = p.parse_args()
    with open(args.input, encoding='utf-8') as f:
        data = json.load(f)
    build_relint(data, args.output)


if __name__ == '__main__':
    main()
