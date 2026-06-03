"""
Geração do Relatório Analítico de Área em formato DOCX.

Replica fielmente a estrutura do modelo CompStat (Anexo do briefing):
  1. Identificação da Área
  2. Indicadores do Período
  3. Distribuição por Tipo de Ocorrência
  4. Análise Temporal (tabela hora × dia)
  5. Perguntas Norteadoras (4 questões respondidas)
  6. Dinâmica Criminal
  7. Fatores de Incidência Criminal (por órgão)
  8. Câmeras identificadas
  9. Painel de Coincidências (Bingo Score resumido)
 10. Plano de Ação e Responsabilização
"""

from __future__ import annotations

import io
from datetime import date
from typing import Optional

import polars as pl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pipeline.bingo import BingoSegmento


# ---------------------------------------------------------------------------
# Helpers de formatação
# ---------------------------------------------------------------------------


def _set_cell_bg(cell, hex_color: str):
    """Define cor de fundo de célula (shading)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_header_row(table, hex_bg: str = "1B3A6B", font_color: str = "FFFFFF"):
    """Aplica estilo de cabeçalho (fundo escuro, texto branco, negrito)."""
    row = table.rows[0]
    for cell in row.cells:
        _set_cell_bg(cell, hex_bg)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(font_color)
                run.font.size = Pt(9)


def _add_section_title(doc: Document, title: str):
    """Adiciona título de seção com formatação CompStat."""
    p = doc.add_paragraph(title)
    p.style = "Heading 2"
    run = p.runs[0] if p.runs else p.add_run(title)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    return p


def _add_table_row(table, cells: list[str], bold_first: bool = False):
    row = table.add_row()
    for i, (cell_obj, val) in enumerate(zip(row.cells, cells)):
        para = cell_obj.paragraphs[0]
        run = para.add_run(str(val) if val is not None else "–")
        run.font.size = Pt(9)
        if bold_first and i == 0:
            run.font.bold = True


def _prioridade_color(prioridade: str) -> str:
    return {
        "CRITICA": "FF4444",
        "ALTA": "FF8800",
        "MEDIA": "FFD700",
        "BAIXA": "AADDAA",
        "URGENTE": "FF4444",
        "PRIORITÁRIA": "FF8800",
        "RECOMENDADA": "FFD700",
    }.get(prioridade.upper(), "FFFFFF")


# ---------------------------------------------------------------------------
# Gerador principal
# ---------------------------------------------------------------------------


def gerar_relatorio(
    area_name: str,
    # Contexto
    aisp: str = "–",
    bpm: str = "–",
    subprefeitura: str = "–",
    base_fm: str = "–",
    bairro: str = "–",
    dp: str = "–",
    orcrim: str = "Não identificada",
    n_cameras: int = 0,
    # Dados quantitativos
    crimes_total: int = 0,
    roubos_total: int = 0,
    furtos_total: int = 0,
    ranking_area: str = "–",
    df_tipos: Optional[pl.DataFrame] = None,  # ranking_tipos_crime()
    df_heatmap: Optional[pl.DataFrame] = None,  # heatmap_hora_dia()
    hora_pico: Optional[int] = None,
    dia_pico: Optional[int] = None,
    # Análise qualitativa (LLM)
    narrativa_dinamica: str = "",
    perguntas_norteadoras: Optional[dict] = None,
    plano_acao: Optional[list[dict]] = None,
    # Bingo Score
    bingos: Optional[list[BingoSegmento]] = None,
    # Fatores
    df_fatores: Optional[pl.DataFrame] = None,
    # Datas
    periodo_crimes: str = "2023–2024",
    periodo_fatores: str = "Fevereiro a Maio de 2026",
) -> bytes:
    """
    Gera e retorna o relatório em bytes (.docx).
    """
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---------------------------------------------------------------------------
    # Cabeçalho
    # ---------------------------------------------------------------------------
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run("RELATÓRIO ANALÍTICO DE ÁREA")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Fundo azul no parágrafo de título (via XML)
    pPr = titulo._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1B3A6B")
    pPr.append(shd)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Subsídio para Reunião de CompStat").font.size = Pt(10)

    doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # Tabela de identificação
    # ---------------------------------------------------------------------------
    _add_section_title(doc, "IDENTIFICAÇÃO DA ÁREA")
    id_table = doc.add_table(rows=5, cols=4)
    id_table.style = "Table Grid"
    rows_data = [
        ("Área FM", area_name, "Número de trechos críticos", str(len(bingos or []))),
        ("AISP", aisp, "Base FM", base_fm),
        ("Bairro", bairro, "Subprefeitura", subprefeitura),
        ("DP", dp, "BPM", bpm),
        (
            "Área sob influência de grupo criminoso",
            orcrim,
            "Câmeras identificadas",
            str(n_cameras),
        ),
    ]
    for r_i, (k1, v1, k2, v2) in enumerate(rows_data):
        row = id_table.rows[r_i]
        for ci, val in enumerate([k1, v1, k2, v2]):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if ci in (0, 2):
                run.font.bold = True
                _set_cell_bg(cell, "E8EEF4")

    doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # 1. Ocorrências Criminais — Indicadores do Período
    # ---------------------------------------------------------------------------
    _add_section_title(doc, "1. OCORRÊNCIAS CRIMINAIS")
    doc.add_paragraph("Dados criminais: " + periodo_crimes).runs[0].font.size = Pt(9)

    ind_table = doc.add_table(rows=2, cols=6)
    ind_table.style = "Table Grid"
    headers = [
        "Período",
        "Roubos",
        "Furtos",
        "Total",
        "Ranking (% áreas FM)",
        "Variação s/ anterior",
    ]
    row0 = ind_table.rows[0]
    for ci, h in enumerate(headers):
        cell = row0.cells[ci]
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(8)
        _set_cell_bg(cell, "1B3A6B")
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    row1 = ind_table.rows[1]
    vals = [
        periodo_crimes,
        str(roubos_total),
        str(furtos_total),
        str(crimes_total),
        ranking_area,
        "–",
    ]
    for ci, v in enumerate(vals):
        run = row1.cells[ci].paragraphs[0].add_run(v)
        run.font.size = Pt(9)

    doc.add_paragraph()

    # Distribuição por tipo
    if df_tipos is not None and not df_tipos.is_empty():
        _add_section_title(doc, "DISTRIBUIÇÃO DE OCORRÊNCIAS POR TIPO")
        tipo_table = doc.add_table(rows=1, cols=4)
        tipo_table.style = "Table Grid"
        for ci, h in enumerate(
            ["Ranking", "Tipo de Ocorrência", "Qtd no período", "%"]
        ):
            cell = tipo_table.rows[0].cells[ci]
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(8)
            _set_cell_bg(cell, "1B3A6B")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for row_data in df_tipos.iter_rows(named=True):
            _add_table_row(
                tipo_table,
                [
                    str(row_data.get("ranking", "–")),
                    str(row_data.get("desc_delito", "–")),
                    str(row_data.get("qtd", "–")),
                    f"{row_data.get('pct', 0):.1f}%",
                ],
            )
        doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # 2. Análise Temporal
    # ---------------------------------------------------------------------------
    _add_section_title(doc, "ANÁLISE TEMPORAL")
    dia_labels = ["Dom", "Seg", "Ter", "Qua", "Qui", "Sex", "Sáb"]
    hora_pico_str = f"{hora_pico:02d}h" if hora_pico is not None else "–"
    dia_pico_str = (
        dia_labels[dia_pico] if dia_pico is not None and 0 <= dia_pico <= 6 else "–"
    )

    if df_heatmap is not None and not df_heatmap.is_empty():
        # Monta matriz dia × hora
        matrix: dict[int, dict[int, int]] = {d: {} for d in range(7)}
        for r_row in df_heatmap.iter_rows(named=True):
            d = r_row.get("dia_semana")
            h = r_row.get("hora")
            c = r_row.get("count", 0)
            if d is not None and h is not None:
                try:
                    matrix[int(d)][int(h)] = int(c)
                except (ValueError, TypeError):
                    pass

        horas_com_dados = sorted({h for d in matrix.values() for h in d})
        if horas_com_dados:
            heat_table = doc.add_table(rows=8, cols=len(horas_com_dados) + 1)
            heat_table.style = "Table Grid"
            # Header horas
            heat_table.rows[0].cells[0].paragraphs[0].add_run(
                "Dia\\Hora"
            ).font.size = Pt(7)
            _set_cell_bg(heat_table.rows[0].cells[0], "E8EEF4")
            for ci, h in enumerate(horas_com_dados, start=1):
                cell = heat_table.rows[0].cells[ci]
                run = cell.paragraphs[0].add_run(f"{h:02d}h")
                run.font.size = Pt(7)
                _set_cell_bg(cell, "E8EEF4")
            # Rows por dia
            for di, dia_label in enumerate(dia_labels):
                row = heat_table.rows[di + 1]
                cell0 = row.cells[0]
                run = cell0.paragraphs[0].add_run(dia_label)
                run.font.size = Pt(7)
                run.font.bold = True
                _set_cell_bg(cell0, "E8EEF4")
                for ci, h in enumerate(horas_com_dados, start=1):
                    cnt = matrix.get(di, {}).get(h, 0)
                    cell = row.cells[ci]
                    run = cell.paragraphs[0].add_run(str(cnt) if cnt > 0 else "–")
                    run.font.size = Pt(7)
                    if cnt >= 10:
                        _set_cell_bg(cell, "FF6666")
                    elif cnt >= 5:
                        _set_cell_bg(cell, "FFAAAA")
                    elif cnt >= 2:
                        _set_cell_bg(cell, "FFE0E0")
            doc.add_paragraph()

    # Resumo temporal
    resumo_table = doc.add_table(rows=2, cols=4)
    resumo_table.style = "Table Grid"
    for ci, h in enumerate(["Período Predominante", "", "Dia / Horário Crítico", ""]):
        cell = resumo_table.rows[0].cells[ci]
        if h:
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(9)
            _set_cell_bg(cell, "1B3A6B")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1 = resumo_table.rows[1]
    r1.cells[0].paragraphs[0].add_run("Período Predominante").font.size = Pt(9)
    r1.cells[1].paragraphs[0].add_run(
        f"Pico às {hora_pico_str} em todos os dias da semana."
    ).font.size = Pt(9)
    r1.cells[2].paragraphs[0].add_run("Dia / Horário Crítico").font.size = Pt(9)
    r1.cells[3].paragraphs[0].add_run(
        f"{dia_pico_str}, às {hora_pico_str}"
    ).font.size = Pt(9)
    doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # 3. Perguntas Norteadoras
    # ---------------------------------------------------------------------------
    if perguntas_norteadoras:
        _add_section_title(doc, "RESUMO EXECUTIVO — PERGUNTAS NORTEADORAS")
        perg_labels = {
            "rota_fm": "Locais de maior incidência criminal estão coincidindo com a rota da FM?",
            "horario_cobertura": "Horário de maior incidência criminal está coincidindo com o QMD da FM?",
            "modelo_emprego": "A dinâmica criminal coincide com o modelo de emprego da FM?",
            "fatores_urbanos": "Os fatores urbanos relevantes estão sendo resolvidos pelos órgãos?",
        }
        perg_table = doc.add_table(rows=1, cols=4)
        perg_table.style = "Table Grid"
        for ci, h in enumerate(
            [
                "Pergunta Norteadora",
                "Diagnóstico com base nos dados",
                "Operação FM / Órgãos",
                "Observações",
            ]
        ):
            cell = perg_table.rows[0].cells[ci]
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(8)
            _set_cell_bg(cell, "1B3A6B")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for key, label in perg_labels.items():
            data = perguntas_norteadoras.get(key, {})
            _add_table_row(
                perg_table,
                [
                    label,
                    data.get("diagnostico", "–"),
                    data.get("operacao_fm", "–"),
                    data.get("obs", "–"),
                ],
            )
        doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # 4. Dinâmica Criminal
    # ---------------------------------------------------------------------------
    _add_section_title(doc, "2. DINÂMICA CRIMINAL")
    din_table = doc.add_table(rows=1, cols=1)
    din_table.style = "Table Grid"
    _set_cell_bg(din_table.rows[0].cells[0], "E8EEF4")
    run = (
        din_table.rows[0].cells[0].paragraphs[0].add_run("ANÁLISE DA DINÂMICA CRIMINAL")
    )
    run.font.bold = True
    run.font.size = Pt(10)

    din_text_row = din_table.add_row()
    para = din_text_row.cells[0].paragraphs[0]
    run = para.add_run(narrativa_dinamica or "Análise qualitativa não disponível.")
    run.font.size = Pt(9)
    doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # 5. Painel de Coincidências (Bingo Score)
    # ---------------------------------------------------------------------------
    if bingos:
        _add_section_title(doc, "PAINEL DE COINCIDÊNCIAS — BINGO SCORE")
        bingo_table = doc.add_table(rows=1, cols=6)
        bingo_table.style = "Table Grid"
        for ci, h in enumerate(
            ["Logradouro", "Score /10", "Prioridade", "Crimes", "Fatores", "Órgãos"]
        ):
            cell = bingo_table.rows[0].cells[ci]
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(8)
            _set_cell_bg(cell, "1B3A6B")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for b in bingos[:15]:
            row = bingo_table.add_row()
            vals = [
                b.logradouro,
                str(b.score_total),
                b.prioridade,
                str(b.crimes_count),
                str(len(b.tipos_fatores)),
                ", ".join(b.orgaos_responsaveis),
            ]
            for ci, (cell, val) in enumerate(zip(row.cells, vals)):
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(8)
                if ci == 2:  # coluna prioridade
                    _set_cell_bg(cell, _prioridade_color(val))

        # Detalhe explicações dos top 5
        doc.add_paragraph()
        p_sub = doc.add_paragraph("Detalhamento das coincidências por trecho:")
        p_sub.runs[0].font.bold = True
        p_sub.runs[0].font.size = Pt(9)
        for b in bingos[:5]:
            if b.score_total >= 3:
                p = doc.add_paragraph()
                p.add_run(
                    f"[{b.score_total}/10 — {b.prioridade}] {b.logradouro}"
                ).font.bold = True
                p.runs[0].font.size = Pt(9)
                for ex in b.explicacao:
                    item = doc.add_paragraph(style="List Bullet")
                    item.add_run(ex).font.size = Pt(8)
        doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # 6. Fatores de Incidência Criminal
    # ---------------------------------------------------------------------------
    if df_fatores is not None and not df_fatores.is_empty():
        _add_section_title(doc, "4. FATORES DE INCIDÊNCIA CRIMINAL")
        doc.add_paragraph(
            "Fatores ambientais e urbanos identificados na área de análise."
        ).runs[0].font.size = Pt(9)

        fat_grouped = (
            df_fatores.group_by("tipo_ocorrencia_descricao")
            .agg(
                [
                    pl.col("logradouro").unique().alias("logradouros"),
                    pl.col("orgao_responsavel").first().alias("orgao"),
                    pl.len().alias("n"),
                ]
            )
            .sort("n", descending=True)
        )

        fat_table = doc.add_table(rows=1, cols=3)
        fat_table.style = "Table Grid"
        for ci, h in enumerate(
            ["Fator Identificado", "Descrição / Logradouros", "Responsável"]
        ):
            cell = fat_table.rows[0].cells[ci]
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(8)
            _set_cell_bg(cell, "1B3A6B")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for r_row in fat_grouped.iter_rows(named=True):
            logrs = list(r_row.get("logradouros") or [])[:4]
            logrs_str = "; ".join(logrs)
            if len(list(r_row.get("logradouros") or [])) > 4:
                logrs_str += "…"
            _add_table_row(
                fat_table,
                [
                    str(r_row.get("tipo_ocorrencia_descricao", "–")),
                    logrs_str,
                    str(r_row.get("orgao", "–")),
                ],
                bold_first=True,
            )
        doc.add_paragraph()

    # ---------------------------------------------------------------------------
    # 7. Plano de Ação e Responsabilização
    # ---------------------------------------------------------------------------
    if plano_acao:
        _add_section_title(doc, "5. PLANO DE AÇÃO E RESPONSABILIZAÇÃO")
        doc.add_paragraph(
            "Esta seção é pré-populada pela IA com base na análise. "
            "Deve ser validada e formalizada durante a reunião de CompStat."
        ).runs[0].font.size = Pt(8)

        plano_table = doc.add_table(rows=1, cols=5)
        plano_table.style = "Table Grid"
        for ci, h in enumerate(
            ["Órgão", "Ação Acordada", "Local Específico", "Justificativa", "Prazo"]
        ):
            cell = plano_table.rows[0].cells[ci]
            run = cell.paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(8)
            _set_cell_bg(cell, "1B3A6B")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for item in plano_acao:
            row = plano_table.add_row()
            vals = [
                item.get("orgao", "–"),
                item.get("acao", "–"),
                item.get("local", "–"),
                item.get("justificativa", "–"),
                item.get("prazo_sugerido", "–"),
            ]
            prioridade_item = item.get("prioridade", "")
            for ci, (cell, val) in enumerate(zip(row.cells, vals)):
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(8)
                if ci == 0:  # orgao
                    run.font.bold = True
                    _set_cell_bg(
                        cell,
                        _prioridade_color(prioridade_item) + "44"
                        if prioridade_item
                        else "F5F5F5",
                    )

        doc.add_paragraph()

    # Rodapé
    doc.add_paragraph()
    rodape = doc.add_paragraph(
        f"CompStat Municipal | Prefeitura do Rio de Janeiro | "
        f"Gerado em {date.today().strftime('%d/%m/%Y')}"
    )
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rodape.runs[0].font.size = Pt(8)
    rodape.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Serializa para bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
